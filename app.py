from flask import Flask, request, send_file, jsonify, redirect, Response
from flask_cors import CORS
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os
import re
import io
import json
import base64
import requests
import time
from datetime import datetime, timedelta

app = Flask(__name__)
# Permitir bodies grandes para upload-attachments (varios PDFs en base64)
app.config['MAX_CONTENT_LENGTH'] = 55 * 1024 * 1024  # 55 MB
# CORS con restricciones de seguridad - solo permitir orígenes específicos
allowed_origins = os.getenv('ALLOWED_ORIGINS', 'https://generador-hojas-vida.web.app,https://generador-hojas-vida.firebaseapp.com').split(',')
CORS(app, origins=allowed_origins, methods=['GET', 'POST'], allow_headers=['Content-Type'])

# Configuración de APIs de iLovePDF
# API Principal: Usada en la API de cursos-certificados (GitHub) - ~250 conversiones
# API de Respaldo: Para cuando se agoten los créditos de la principal - ~250 conversiones
# Total disponible: ~500 conversiones
ILOVEPDF_APIS = [
    {
        'name': 'primary',
        # API Principal - Credenciales de la API de cursos-certificados
        # TODO: Reemplazar con las credenciales reales de la API principal
        'public_key': os.getenv('ILOVEPDF_PRIMARY_PUBLIC_KEY', None),  # Agregar aquí la public_key principal
        'secret_key': os.getenv('ILOVEPDF_PRIMARY_SECRET_KEY', None)   # Agregar aquí la secret_key principal
    },
    {
        'name': 'backup',
        # API de Respaldo - Se usa automáticamente cuando la principal se queda sin créditos
        # NOTA: Estas credenciales deben moverse a variables de entorno en producción
        'public_key': os.getenv('ILOVEPDF_BACKUP_PUBLIC_KEY', 'project_public_e8de4c9dde8d3130930dc8f9620f9fd0_4gcUq34631a35630e89502c9cb2229d123ff4'),
        'secret_key': os.getenv('ILOVEPDF_BACKUP_SECRET_KEY', 'secret_key_5f1ab1bb9dc866aadc8a05671e460491_zNqoaf28f8b33e1755f025940359d1d4a70a3')
    }
]

# INSTRUCCIONES PARA CONFIGURAR LA API PRINCIPAL:
# 1. Busca las credenciales de iLovePDF en el código de la API de cursos-certificados (GitHub)
# 2. Reemplaza los valores None arriba con las credenciales reales, O
# 3. Configura variables de entorno en Render:
#    - ILOVEPDF_PRIMARY_PUBLIC_KEY = "tu_public_key_aqui"
#    - ILOVEPDF_PRIMARY_SECRET_KEY = "tu_secret_key_aqui"

# Variable para rastrear qué API está activa
current_api_index = 0

# Meses en español
MESES = {
    1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
    5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
    9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
}

def formatear_fecha(fecha_str: str) -> str:
    """Convierte una fecha a formato '20 de noviembre de 1990'
    Acepta formatos: YYYY-MM-DD, DD/MM/YYYY, DD-MM-YYYY
    """
    if not fecha_str:
        return ''
    
    fecha_str = fecha_str.strip()
    
    try:
        # Intentar formato YYYY-MM-DD (estándar de date picker)
        if fecha_str.count('-') == 2:
            partes = fecha_str.split('-')
            if len(partes[0]) == 4:  # YYYY-MM-DD
                año, mes, dia = partes
                año, mes, dia = int(año), int(mes), int(dia)
                if 1 <= mes <= 12 and 1 <= dia <= 31:
                    return f"{dia} de {MESES[mes]} de {año}"
        
        # Intentar formato DD/MM/YYYY o DD-MM-YYYY
        if fecha_str.count('/') == 2 or fecha_str.count('-') == 2:
            separador = '/' if '/' in fecha_str else '-'
            partes = fecha_str.split(separador)
            if len(partes) == 3:
                # Determinar si es DD/MM/YYYY o MM/DD/YYYY
                # Asumimos DD/MM/YYYY si el primer número es <= 31
                if int(partes[0]) <= 31:
                    dia, mes, año = int(partes[0]), int(partes[1]), int(partes[2])
                else:
                    mes, dia, año = int(partes[0]), int(partes[1]), int(partes[2])
                
                if 1 <= mes <= 12 and 1 <= dia <= 31:
                    return f"{dia} de {MESES[mes]} de {año}"
    except (ValueError, IndexError, KeyError):
        pass
    
    # Si no se puede parsear, devolver la fecha original
    return fecha_str

@app.route('/', methods=['GET'])
def root():
    """Endpoint raíz para verificar que el servidor está funcionando"""
    return jsonify({
        "status": "ok", 
        "message": "API de Generación de Hojas de Vida funcionando",
        "endpoints": {
            "/health": "GET - Verificar estado del servidor",
            "/upload-attachments": "POST - Subir anexos a Google Drive o Firebase (clientName, clientId, attachments). Primero Drive; no requiere activar Firebase Storage.",
            "/drive-download": "GET - Descargar archivo (file_id, file_name). Firebase: anexos/...; Drive: ID de archivo. Para admin/solicitudes (Fabian).",
            "/list-folder": "GET - Listar archivos de carpeta (folder_id). Firebase: anexos/...; Drive: ID de carpeta. Para admin/solicitudes (Fabian).",
            "/generate-word": "POST - Generar documento Word (Hoja de Vida)",
            "/generate-cuenta-cobro": "POST - Generar cuenta de cobro desde template",
            "/convert-word-to-pdf": "POST - Convertir Word a PDF usando iLovePDF"
        }
    })

@app.route('/health', methods=['GET'])
def health():
    """Endpoint de salud para verificar que el servidor está funcionando"""
    return jsonify({"status": "ok", "message": "API funcionando correctamente"})

@app.errorhandler(413)
def request_entity_too_large(e):
    """Cuerpo de la petición demasiado grande (varios PDFs en base64). Pedir comprimir archivos."""
    return jsonify({
        'error': 'Los archivos son demasiado grandes. Comprime los PDFs (menos de 2 MB cada uno) o sube menos archivos a la vez.',
        'success': False
    }), 413

# --- Subida de anexos: primero Google Drive (sin activar Firebase Storage); si no está configurado, fallback a Firebase ---
ATTACHMENT_NAMES = {
    'cedula': 'Cedula',
    'actaBachiller': 'Acta_Bachiller',
    'diplomaBachiller': 'Diploma_Bachiller',
    'actaOtro': 'Acta_Otro_Estudio',
    'diplomaOtro': 'Diploma_Otro_Estudio',
    'cursos': 'Cursos',
    'antecedentes': 'Antecedentes',
    'rut': 'RUT',
    'pension': 'Certificado_Pension',
    'eps': 'Certificado_EPS',
    'vacunas': 'Vacunas',
    'arl': 'Certificado_ARL',
    'otros': 'Otros_Documentos',
    'contrato': 'Contrato',
}

_firebase_app = None
_drive_service = None

def get_drive_service():
    """Devuelve el cliente de Google Drive API si está configurado (GOOGLE_DRIVE_SERVICE_ACCOUNT_JSON). Así no hace falta activar Firebase Storage."""
    global _drive_service
    if _drive_service is not None:
        return _drive_service
    cred_json = os.getenv('GOOGLE_DRIVE_SERVICE_ACCOUNT_JSON')
    if not cred_json:
        return None
    try:
        from google.oauth2.service_account import Credentials
        from googleapiclient.discovery import build
        cred_dict = json.loads(cred_json)
        creds = Credentials.from_service_account_info(cred_dict, scopes=['https://www.googleapis.com/auth/drive.file'])
        _drive_service = build('drive', 'v3', credentials=creds)
        return _drive_service
    except Exception as e:
        print('Drive init error:', e)
        _drive_service = None
        return None

def get_firebase_bucket():
    """Inicializa Firebase Admin (una vez) y devuelve el bucket de Storage."""
    global _firebase_app
    if _firebase_app is not None:
        from firebase_admin import storage
        return storage.bucket()
    cred_json = os.getenv('FIREBASE_SERVICE_ACCOUNT_JSON')
    if not cred_json:
        return None
    try:
        import firebase_admin
        from firebase_admin import credentials, storage
        cred_dict = json.loads(cred_json)
        _firebase_app = firebase_admin.initialize_app(
            credentials.Certificate(cred_dict),
            {'storageBucket': os.getenv('FIREBASE_STORAGE_BUCKET', 'generador-hojas-vida.appspot.com')},
        )
        return storage.bucket()
    except Exception as e:
        print('Firebase init error:', e)
        _firebase_app = None
        return None

def _data_url_to_bytes(data_url):
    if not data_url or not data_url.startswith('data:'):
        return None, 'application/octet-stream'
    comma = data_url.find(',')
    b64 = data_url[comma + 1:] if comma >= 0 else ''
    mime = 'application/octet-stream'
    if data_url[:comma]:
        part = data_url[:comma].replace('data:', '').split(';')[0].strip()
        if part:
            mime = part
    try:
        return base64.b64decode(b64), mime
    except Exception:
        return None, mime

def _upload_attachments_to_drive(client_name, client_id, attachments):
    """Sube anexos a Google Drive. Misma forma de respuesta que Firebase para el frontend."""
    drive = get_drive_service()
    if not drive:
        return None
    name = re.sub(r'[\s/\\?*:]+', '_', client_name).strip('_') or 'Cliente'
    cid = re.sub(r'[\s/\\?*:]+', '_', client_id).strip('_') or ''
    folder_name = f"{name}_{cid}".replace('__', '_').strip('_') or 'cliente_doc'
    parent_id = os.getenv('GOOGLE_DRIVE_FOLDER_ID', '').strip() or None
    try:
        from googleapiclient.http import MediaIoBaseUpload
        # Crear carpeta por cliente
        folder_meta = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
        if parent_id:
            folder_meta['parents'] = [parent_id]
        folder = drive.files().create(body=folder_meta, fields='id,webViewLink').execute()
        folder_id = folder.get('id')
        folder_link = folder.get('webViewLink') or f'https://drive.google.com/drive/folders/{folder_id}'
        uploaded_files = []
        errors = []
        for key, att in attachments.items():
            if not att or not att.get('dataUrl'):
                continue
            label = ATTACHMENT_NAMES.get(key, key)
            fname = (att.get('name') or 'documento').strip()
            ext = fname.split('.')[-1] if '.' in fname else 'pdf'
            if ext.lower() not in ('pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx'):
                ext = 'pdf'
            file_name = f'{label}.{ext}'
            raw, content_type = _data_url_to_bytes(att['dataUrl'])
            if raw is None:
                errors.append(f'Error decodificando {key} ({file_name})')
                continue
            try:
                media = MediaIoBaseUpload(io.BytesIO(raw), mimetype=content_type or 'application/octet-stream', resumable=True)
                file_meta = {'name': file_name, 'parents': [folder_id]}
                up = drive.files().create(body=file_meta, media_body=media, fields='id,webViewLink').execute()
                web_link = up.get('webViewLink') or f'https://drive.google.com/file/d/{up.get("id")}/view'
                uploaded_files.append({'key': key, 'name': file_name, 'file_id': up.get('id'), 'web_link': web_link})
            except Exception as e:
                err_msg = str(e).split('\n')[0][:200] if e else 'Error desconocido'
                errors.append(f'Error subiendo {key} ({file_name}): {err_msg}')
        return {
            'success': True,
            'folder_name': folder_name,
            'folder_id': folder_id,
            'drive_folder_link': folder_link,
            'uploaded_files': uploaded_files,
            'errors': errors,
            'storage_type': 'drive',
            'message': f'Se subieron {len(uploaded_files)} archivo(s) a Google Drive.',
        }
    except Exception as e:
        print('Drive upload error:', e)
        return None

@app.route('/upload-attachments', methods=['POST', 'OPTIONS'])
def upload_attachments():
    """
    Sube anexos: primero intenta Google Drive (no requiere activar Firebase Storage).
    Si Drive no está configurado, usa Firebase Storage. Mismo cuerpo y respuesta para el frontend.
    """
    if request.method == 'OPTIONS':
        return '', 204
    data = request.get_json() or {}
    client_name = (data.get('clientName') or '').strip()
    client_id = (data.get('clientId') or '').strip()
    attachments = data.get('attachments') or {}
    if not client_name or not client_id:
        return jsonify({'error': 'Se requiere clientName y clientId', 'success': False}), 400
    if not attachments:
        return jsonify({'error': 'No se proporcionaron anexos (attachments)', 'success': False}), 400
    name = re.sub(r'[\s/\\?*:]+', '_', client_name).strip('_') or 'Cliente'
    cid = re.sub(r'[\s/\\?*:]+', '_', client_id).strip('_') or ''
    folder_name = f"{name}_{cid}".replace('__', '_').strip('_') or 'cliente_doc'
    folder_path = f'anexos/{folder_name}'

    # 1) Intentar Google Drive primero (no hace falta activar Firebase Storage)
    result = _upload_attachments_to_drive(client_name, client_id, attachments)
    if result is not None:
        return jsonify(result)

    # 2) Fallback a Firebase Storage si está configurado
    bucket = get_firebase_bucket()
    if not bucket:
        return jsonify({
            'error': 'Configura Google Drive en Render (GOOGLE_DRIVE_SERVICE_ACCOUNT_JSON y opcional GOOGLE_DRIVE_FOLDER_ID). Así no necesitas activar Firebase Storage. Ver GOOGLE_DRIVE_SETUP.md.',
            'success': False,
        }), 503
    uploaded_files = []
    errors = []
    try:
        from firebase_admin import storage
        for key, att in attachments.items():
            if not att or not att.get('dataUrl'):
                continue
            label = ATTACHMENT_NAMES.get(key, key)
            fname = (att.get('name') or 'documento').strip()
            ext = fname.split('.')[-1] if '.' in fname else 'pdf'
            if ext.lower() not in ('pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx'):
                ext = 'pdf'
            file_name = f'{label}.{ext}'
            file_path = f'{folder_path}/{file_name}'
            raw, content_type = _data_url_to_bytes(att['dataUrl'])
            if raw is None:
                errors.append(f'Error decodificando {key} ({file_name})')
                continue
            try:
                blob = bucket.blob(file_path)
                blob.upload_from_string(raw, content_type=content_type or 'application/octet-stream')
            except Exception as e:
                err_msg = str(e).split('\n')[0][:200] if e else 'Error desconocido'
                errors.append(f'Error subiendo {key} ({file_name}): {err_msg}')
                continue
            try:
                url = blob.generate_signed_url(expiration=timedelta(days=365 * 10), method='GET')
            except Exception:
                url = f'https://firebasestorage.googleapis.com/v0/b/{bucket.name}/o/{file_path.replace("/", "%2F")}?alt=media'
            uploaded_files.append({
                'key': key,
                'name': file_name,
                'file_id': file_path,
                'web_link': url,
            })
        first_url = uploaded_files[0]['web_link'] if uploaded_files else ''
        drive_folder_link = (first_url.rsplit('/', 1)[0] + '/') if first_url else folder_path
        return jsonify({
            'success': True,
            'folder_name': folder_name,
            'folder_id': folder_path,
            'drive_folder_link': drive_folder_link,
            'uploaded_files': uploaded_files,
            'errors': errors if errors else [],
            'storage_type': 'firebase',
            'message': f'Se subieron {len(uploaded_files)} archivo(s) a Firebase.',
        })
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500

@app.route('/drive-download', methods=['GET'])
def drive_download():
    """
    Descarga un archivo para el admin (solicitudes / flujo Fabian).
    - file_id: ruta en Firebase (anexos/Nombre_123/Archivo.pdf) o ID de archivo en Google Drive.
    - file_name: nombre para la descarga.
    Así el admin puede Ver/Descargar tanto archivos en Firebase Storage como en Drive.
    """
    file_id = (request.args.get('file_id') or '').strip()
    file_name = (request.args.get('file_name') or 'documento.pdf').strip()
    if not file_id:
        return jsonify({'error': 'Falta file_id'}), 400
    # Sanitizar nombre para Content-Disposition
    safe_name = re.sub(r'[^\w\s\-\.]', '_', file_name)[:200] or 'documento.pdf'
    try:
        # 1) Firebase Storage: file_id = anexos/Nombre_123/Archivo.pdf
        if file_id.startswith('anexos/'):
            bucket = get_firebase_bucket()
            if not bucket:
                return jsonify({'error': 'Firebase Storage no configurado'}), 503
            blob = bucket.blob(file_id)
            data = blob.download_as_bytes()
            return Response(
                data,
                mimetype='application/octet-stream',
                headers={'Content-Disposition': f'attachment; filename="{safe_name}"'}
            )
        # 2) Google Drive: file_id = ID del archivo en Drive
        drive = get_drive_service()
        if not drive:
            return jsonify({'error': 'Google Drive no configurado'}), 503
        content = drive.files().get_media(fileId=file_id).execute()
        return Response(
            content,
            mimetype='application/octet-stream',
            headers={'Content-Disposition': f'attachment; filename="{safe_name}"'}
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/list-folder', methods=['GET'])
def list_folder():
    """
    Lista archivos de una carpeta para el admin (solicitudes / flujo Fabian).
    - folder_id: ruta en Firebase (anexos/Nombre_123) o ID de carpeta en Google Drive.
    Devuelve [{ id, name, public_url? }]. public_url solo en Firebase; en Drive se usa /drive-download.
    """
    folder_id = (request.args.get('folder_id') or '').strip()
    if not folder_id:
        return jsonify({'error': 'Falta folder_id', 'files': []}), 400
    try:
        # 1) Firebase Storage: folder_id = anexos/Nombre_123
        if folder_id.startswith('anexos/'):
            bucket = get_firebase_bucket()
            if not bucket:
                return jsonify({'error': 'Firebase Storage no configurado', 'files': []}), 503
            prefix = folder_id.rstrip('/') + '/'
            blobs = list(bucket.list_blobs(prefix=prefix))
            files = []
            for b in blobs:
                if b.name == prefix or not b.name.startswith(prefix):
                    continue
                name = b.name[len(prefix):].strip()
                if not name:
                    continue
                try:
                    url = b.generate_signed_url(expiration=timedelta(days=365), method='GET')
                except Exception:
                    url = f'https://firebasestorage.googleapis.com/v0/b/{bucket.name}/o/{b.name.replace("/", "%2F")}?alt=media'
                files.append({'id': b.name, 'name': name, 'public_url': url})
            return jsonify({'files': files})
        # 2) Google Drive: folder_id = ID de la carpeta
        drive = get_drive_service()
        if not drive:
            return jsonify({'error': 'Google Drive no configurado', 'files': []}), 503
        q = f"'{folder_id}' in parents and trashed = false"
        results = drive.files().list(q=q, pageSize=100, fields='files(id, name)').execute()
        items = results.get('files', [])
        files = [{'id': f['id'], 'name': f.get('name', ''), 'public_url': None} for f in items]
        return jsonify({'files': files})
    except Exception as e:
        return jsonify({'error': str(e), 'files': []}), 500

def convert_word_to_pdf_with_ilovepdf(word_file_bytes, filename='document.docx'):
    """
    Convierte un archivo Word a PDF usando la API de iLovePDF con fallback automático
    si se acaban los créditos.
    """
    global current_api_index
    
    max_retries = len(ILOVEPDF_APIS)
    
    for attempt in range(max_retries):
        api_config = ILOVEPDF_APIS[current_api_index]
        
        # Si no hay credenciales configuradas, saltar esta API
        if not api_config['public_key'] or not api_config['secret_key']:
            current_api_index = (current_api_index + 1) % len(ILOVEPDF_APIS)
            continue
        
        try:
            # Paso 1: Autenticarse y obtener token
            auth_url = 'https://api.ilovepdf.com/v1/auth'
            auth_response = requests.post(auth_url, json={
                'public_key': api_config['public_key']
            })
            
            # Verificar respuesta de autenticación
            if auth_response.status_code != 200:
                error_text = auth_response.text.lower()
                error_json = {}
                try:
                    error_json = auth_response.json()
                except:
                    pass
                
                # Detectar errores de créditos
                if any(ind in error_text for ind in ['credits', 'quota', 'limit', 'exceeded', 'insufficient', 'balance']):
                    raise Exception(f"CREDITS_EXHAUSTED: {auth_response.text}")
                
                raise Exception(f"Error de autenticación ({auth_response.status_code}): {auth_response.text}")
            
            auth_data = auth_response.json()
            token = auth_data.get('token')
            
            if not token:
                raise Exception("No se recibió token de autenticación")
            
            # Paso 2: Iniciar tarea de conversión
            start_url = 'https://api.ilovepdf.com/v1/start/officepdf'
            headers = {'Authorization': f'Bearer {token}'}
            start_response = requests.get(start_url, headers=headers)
            
            if start_response.status_code != 200:
                error_text = start_response.text.lower()
                if any(ind in error_text for ind in ['credits', 'quota', 'limit', 'exceeded', 'insufficient', 'balance']):
                    raise Exception(f"CREDITS_EXHAUSTED: {start_response.text}")
                raise Exception(f"Error al iniciar tarea ({start_response.status_code}): {start_response.text}")
            
            task_data = start_response.json()
            server = task_data.get('server')
            task = task_data.get('task')
            
            if not server or not task:
                raise Exception("No se recibieron datos de servidor o tarea")
            
            # Paso 3: Subir archivo Word
            upload_url = f'https://{server}/v1/upload'
            files = {'file': (filename, word_file_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            upload_response = requests.post(upload_url, files=files, headers=headers)
            
            if upload_response.status_code != 200:
                error_text = upload_response.text.lower()
                if any(ind in error_text for ind in ['credits', 'quota', 'limit', 'exceeded', 'insufficient', 'balance']):
                    raise Exception(f"CREDITS_EXHAUSTED: {upload_response.text}")
                raise Exception(f"Error al subir archivo ({upload_response.status_code}): {upload_response.text}")
            
            upload_data = upload_response.json()
            server_filename = upload_data.get('server_filename')
            
            if not server_filename:
                raise Exception("No se recibió nombre de archivo del servidor")
            
            # Paso 4: Procesar conversión
            process_url = f'https://{server}/v1/process'
            process_data = {
                'task': task,
                'tool': 'officepdf',
                'files': [{'server_filename': server_filename, 'filename': filename}]
            }
            process_response = requests.post(process_url, json=process_data, headers=headers)
            
            if process_response.status_code != 200:
                error_text = process_response.text.lower()
                if any(ind in error_text for ind in ['credits', 'quota', 'limit', 'exceeded', 'insufficient', 'balance']):
                    raise Exception(f"CREDITS_EXHAUSTED: {process_response.text}")
                raise Exception(f"Error al procesar ({process_response.status_code}): {process_response.text}")
            
            # Esperar un momento para que el procesamiento termine
            time.sleep(1)
            
            # Paso 5: Descargar PDF resultante
            download_url = f'https://{server}/v1/download/{task}'
            download_response = requests.get(download_url, headers=headers)
            
            if download_response.status_code != 200:
                error_text = download_response.text.lower()
                if any(ind in error_text for ind in ['credits', 'quota', 'limit', 'exceeded', 'insufficient', 'balance']):
                    raise Exception(f"CREDITS_EXHAUSTED: {download_response.text}")
                raise Exception(f"Error al descargar PDF ({download_response.status_code}): {download_response.text}")
            
            # Si llegamos aquí, la conversión fue exitosa
            print(f"✅ Conversión exitosa usando API {api_config['name']}")
            return download_response.content
            
        except Exception as e:
            error_message = str(e)
            is_credits_error = 'CREDITS_EXHAUSTED' in error_message or any(
                ind in error_message.lower() for ind in ['credits', 'quota', 'limit', 'exceeded', 'insufficient', 'balance', '401', '403']
            )
            
            if is_credits_error:
                # Cambiar a la siguiente API
                current_api_index = (current_api_index + 1) % len(ILOVEPDF_APIS)
                print(f"⚠️ Créditos agotados en API {api_config['name']}. Cambiando a API de respaldo...")
                
                # Si no hay más APIs, lanzar error
                if attempt == max_retries - 1:
                    raise Exception(f"Todas las APIs de iLovePDF han agotado sus créditos. Último error: {error_message}")
                
                # Continuar con el siguiente intento
                continue
            else:
                # Error diferente, relanzar
                raise e
    
    raise Exception("No se pudo convertir el archivo después de intentar todas las APIs disponibles")

@app.route('/convert-word-to-pdf', methods=['POST'])
def convert_word_to_pdf():
    """
    Convierte un documento Word a PDF usando iLovePDF con fallback automático
    """
    try:
        # Verificar si se envió un archivo
        if 'file' not in request.files:
            return jsonify({"error": "No se proporcionó ningún archivo"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"error": "Nombre de archivo vacío"}), 400
        
        # Leer el archivo
        word_file_bytes = file.read()
        
        # Convertir a PDF
        pdf_bytes = convert_word_to_pdf_with_ilovepdf(word_file_bytes, file.filename)
        
        # Preparar respuesta
        output = io.BytesIO(pdf_bytes)
        output.seek(0)
        
        # Nombre del archivo PDF
        pdf_filename = file.filename.replace('.docx', '.pdf').replace('.doc', '.pdf')
        if not pdf_filename.endswith('.pdf'):
            pdf_filename += '.pdf'
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=pdf_filename
        )
        
    except Exception as e:
        import traceback
        return jsonify({
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500

@app.route('/generate-word', methods=['POST'])
def generate_word():
    """Genera un documento Word desde cero con todos los datos recibidos"""
    try:
        data = request.json
        
        # Obtener datos básicos
        nombre = data.get('fullName', '').strip()
        cedula = data.get('idNumber', '').strip()
        fecha_raw = data.get('birthDate', '').strip()
        fecha = formatear_fecha(fecha_raw)  # Convertir a formato "20 de noviembre de 1990"
        telefono = data.get('phone', '').strip()
        direccion = data.get('address', '').strip()
        ciudad = data.get('place', '').strip()
        estado_civil = data.get('estadoCivil', '').strip().upper()  # Convertir a mayúsculas
        correo = data.get('email', '').strip()
        exp = data.get('idIssuePlace', '').strip()
        texto_perfil = data.get('profile', '').strip()
        
        # Obtener referencias
        referencias_familiares = data.get('referenciasFamiliares', [])
        referencias_personales = data.get('referenciasPersonales', [])
        
        # Obtener experiencias laborales
        experiencias = data.get('experiencias', [])
        
        # Obtener formaciones académicas
        formaciones = data.get('formaciones', [])
        high_school = data.get('highSchool', '').strip()
        institution = data.get('institution', '').strip()
        
        # Crear un nuevo documento desde cero
        doc = Document()
        
        # Configurar encabezado con fondo azul
        section = doc.sections[0]
        header = section.header
        
        # Limpiar párrafos existentes del encabezado
        for para in header.paragraphs:
            para.clear()
        
        # Crear nuevo párrafo para el encabezado
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Agregar fondo azul al encabezado usando XML
        header_xml = header_para._element
        pPr = header_xml.get_or_add_pPr()
        
        # Crear elemento de sombreado (fondo azul #5B9BD5)
        shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="5B9BD5" w:val="clear"/>')
        pPr.append(shd)
        
        # Agregar espaciado superior e inferior para que el fondo se vea como una barra
        spacing = parse_xml(r'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:before="240" w:after="240"/>')
        pPr.append(spacing)
        
        # Agregar indentación derecha para que el texto no toque el borde
        ind = parse_xml(r'<w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:right="360"/>')
        pPr.append(ind)
        
        # Agregar el nombre en blanco, negrita
        header_run = header_para.add_run(nombre.upper())
        header_run.font.name = "Calibri"
        header_run.font.size = Pt(11)
        header_run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
        header_run.bold = True
        
        # Agregar espacio
        doc.add_paragraph()
        
        # Nombre principal (Cambria 18, color #4472C4, mayúsculas, negrita)
        p_nombre = doc.add_paragraph()
        run_nombre = p_nombre.add_run(nombre.upper())
        run_nombre.font.name = "Cambria"
        run_nombre.font.size = Pt(18)
        run_nombre.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_nombre.bold = True
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        
        # Información personal - etiquetas en negrita azul, valores en negro
        p_cedula = doc.add_paragraph()
        run_cedula_label = p_cedula.add_run("Número de cédula: ")
        run_cedula_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_cedula_label.bold = True
        run_cedula_valor = p_cedula.add_run(cedula)
        run_cedula_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        p_fecha = doc.add_paragraph()
        run_fecha_label = p_fecha.add_run("Fecha de nacimiento: ")
        run_fecha_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_fecha_label.bold = True
        run_fecha_valor = p_fecha.add_run(fecha)
        run_fecha_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        p_tel = doc.add_paragraph()
        run_tel_label = p_tel.add_run("Teléfono móvil: ")
        run_tel_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_tel_label.bold = True
        run_tel_valor = p_tel.add_run(telefono)
        run_tel_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        p_dir = doc.add_paragraph()
        run_dir_label = p_dir.add_run("Dirección: ")
        run_dir_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_dir_label.bold = True
        run_dir_valor = p_dir.add_run(direccion)
        run_dir_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        p_ciu = doc.add_paragraph()
        run_ciu_label = p_ciu.add_run("Ciudad: ")
        run_ciu_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_ciu_label.bold = True
        run_ciu_valor = p_ciu.add_run(ciudad)
        run_ciu_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        p_est = doc.add_paragraph()
        run_est_label = p_est.add_run("Estado civil: ")
        run_est_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_est_label.bold = True
        run_est_valor = p_est.add_run(estado_civil)
        run_est_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        if correo:
            p_corr = doc.add_paragraph()
            run_corr_label = p_corr.add_run("Correo: ")
            run_corr_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            run_corr_label.bold = True
            run_corr_valor = p_corr.add_run(correo)
            run_corr_valor.font.color.rgb = RGBColor(0, 0, 0)
        
        # Perfil Profesional - título en azul, negrita, mayúsculas
        if texto_perfil:
            p_perfil_titulo = doc.add_paragraph()
            p_perfil_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_perfil_titulo = p_perfil_titulo.add_run("PERFIL PROFESIONAL")
            run_perfil_titulo.bold = True
            run_perfil_titulo.font.size = Pt(12)
            run_perfil_titulo.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            doc.add_paragraph()
            
            p_perfil_texto = doc.add_paragraph()
            p_perfil_texto.add_run(texto_perfil)
        
        # Si NO hay experiencia laboral, agregar formación académica en la hoja 1
        if not experiencias:
            # Solo agregar formación académica si hay datos
            if high_school or institution or formaciones:
                doc.add_paragraph()
                doc.add_paragraph()
                
                # Formación Académica - título en azul, negrita, mayúsculas (hoja 1)
                p_formacion_titulo = doc.add_paragraph()
                p_formacion_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_formacion_titulo = p_formacion_titulo.add_run("FORMACIÓN ACADÉMICA")
                run_formacion_titulo.bold = True
                run_formacion_titulo.font.size = Pt(12)
                run_formacion_titulo.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                doc.add_paragraph()
                
                # Formación académica sin tabla, solo texto alineado
                if high_school or institution:
                    p_sec = doc.add_paragraph()
                    run_sec_label = p_sec.add_run("BACHILLER: ")
                    run_sec_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_sec_label.bold = True
                    run_sec_valor = p_sec.add_run(high_school)
                    run_sec_valor.font.color.rgb = RGBColor(0, 0, 0)
                    
                    p_inst = doc.add_paragraph()
                    run_inst_label = p_inst.add_run("INSTITUCION: ")
                    run_inst_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_inst_label.bold = True
                    run_inst_valor = p_inst.add_run(institution)
                    run_inst_valor.font.color.rgb = RGBColor(0, 0, 0)
                
                # Formación técnica/universitaria (puede haber múltiples)
                # Solo agregar si NO es "Bachiller" (ya está arriba)
                for form in formaciones:
                    tipo_form = form.get('tipo', '').strip().upper()
                    nombre_form = form.get('nombre', '').strip()
                    
                    # Filtrar: no mostrar si es "BACHILLER" (ya está en la sección de secundaria)
                    if tipo_form and tipo_form != 'BACHILLER' and nombre_form:
                        doc.add_paragraph()
                        p_tec = doc.add_paragraph()
                        run_tec_label = p_tec.add_run(tipo_form)
                        run_tec_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_tec_label.bold = True
                        # El valor en la misma línea con dos puntos
                        run_tec_colon = p_tec.add_run(": ")
                        run_tec_colon.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_tec_colon.bold = True
                        run_tec_valor = p_tec.add_run(nombre_form)
                        run_tec_valor.font.color.rgb = RGBColor(0, 0, 0)
                
                # Salto de página después de formación académica (inicio de hoja 2 para referencias)
                p_break1 = doc.add_paragraph()
                run_break1 = p_break1.add_run()
                run_break1.add_break(WD_BREAK.PAGE)
        else:
            # Si hay experiencia laboral, salto de página después del perfil profesional (inicio de hoja 2)
            p_break1 = doc.add_paragraph()
            run_break1 = p_break1.add_run()
            run_break1.add_break(WD_BREAK.PAGE)
            
            # Formación Académica - título en azul, negrita, mayúsculas (hoja 2)
            # Solo agregar si hay datos
            if high_school or institution or formaciones:
                p_formacion_titulo = doc.add_paragraph()
                p_formacion_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_formacion_titulo = p_formacion_titulo.add_run("FORMACIÓN ACADÉMICA")
                run_formacion_titulo.bold = True
                run_formacion_titulo.font.size = Pt(12)
                run_formacion_titulo.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                doc.add_paragraph()
                
                # Formación académica sin tabla, solo texto alineado
                if high_school or institution:
                    p_sec = doc.add_paragraph()
                    run_sec_label = p_sec.add_run("BACHILLER: ")
                    run_sec_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_sec_label.bold = True
                    run_sec_valor = p_sec.add_run(high_school)
                    run_sec_valor.font.color.rgb = RGBColor(0, 0, 0)
                    
                    p_inst = doc.add_paragraph()
                    run_inst_label = p_inst.add_run("INSTITUCION: ")
                    run_inst_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_inst_label.bold = True
                    run_inst_valor = p_inst.add_run(institution)
                    run_inst_valor.font.color.rgb = RGBColor(0, 0, 0)
                
                # Formación técnica/universitaria (puede haber múltiples)
                # Solo agregar si NO es "Bachiller" (ya está arriba)
                for form in formaciones:
                    tipo_form = form.get('tipo', '').strip().upper()
                    nombre_form = form.get('nombre', '').strip()
                    
                    # Filtrar: no mostrar si es "BACHILLER" (ya está en la sección de secundaria)
                    if tipo_form and tipo_form != 'BACHILLER' and nombre_form:
                        doc.add_paragraph()
                        p_tec = doc.add_paragraph()
                        run_tec_label = p_tec.add_run(tipo_form)
                        run_tec_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_tec_label.bold = True
                        # El valor en la misma línea con dos puntos
                        run_tec_colon = p_tec.add_run(": ")
                        run_tec_colon.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_tec_colon.bold = True
                        run_tec_valor = p_tec.add_run(nombre_form)
                        run_tec_valor.font.color.rgb = RGBColor(0, 0, 0)
                
                doc.add_paragraph()
                doc.add_paragraph()
        
        # Experiencia Laboral - título en azul, negrita, mayúsculas, centrado (hoja 2, solo si hay experiencia)
        if experiencias:
            p_exp_titulo = doc.add_paragraph()
            p_exp_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_exp_titulo = p_exp_titulo.add_run("EXPERIENCIA LABORAL")
            run_exp_titulo.bold = True
            run_exp_titulo.font.size = Pt(12)
            run_exp_titulo.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            doc.add_paragraph()
            
            for experiencia in experiencias:
                # Las experiencias pueden venir con 'empresa' o 'local', 'cargo' o 'cargo', 'tiempo' o 'fechaInicio/fechaFin'
                empresa = experiencia.get('empresa', experiencia.get('local', '')).strip()
                cargo = experiencia.get('cargo', '').strip()
                tiempo = experiencia.get('tiempo', '')
                
                # Si no viene tiempo, construirlo desde fechaInicio y fechaFin
                if not tiempo:
                    fecha_inicio = experiencia.get('fechaInicio', '').strip()
                    fecha_fin = experiencia.get('fechaFin', '').strip()
                    if fecha_inicio and fecha_fin:
                        tiempo = f"Desde {fecha_inicio} hasta {fecha_fin}"
                
                if empresa and cargo:
                    p_estab = doc.add_paragraph()
                    run_estab_label = p_estab.add_run("ESTABLECIMIENTO: ")
                    run_estab_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_estab_label.bold = True
                    run_estab_valor = p_estab.add_run(empresa)
                    run_estab_valor.font.color.rgb = RGBColor(0, 0, 0)
                    
                    p_cargo = doc.add_paragraph()
                    run_cargo_label = p_cargo.add_run("CARGO: ")
                    run_cargo_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_cargo_label.bold = True
                    run_cargo_valor = p_cargo.add_run(cargo)
                    run_cargo_valor.font.color.rgb = RGBColor(0, 0, 0)
                    
                    if tiempo:
                        p_periodo = doc.add_paragraph()
                        run_periodo_label = p_periodo.add_run("PERIODO LABORAL: ")
                        run_periodo_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_periodo_label.bold = True
                        run_periodo_valor = p_periodo.add_run(tiempo)
                        run_periodo_valor.font.color.rgb = RGBColor(0, 0, 0)
                    
                    doc.add_paragraph()
                    doc.add_paragraph()
            
            # Si hay experiencia, salto de página para referencias (hoja 3)
            p_break2 = doc.add_paragraph()
            run_break2 = p_break2.add_run()
            run_break2.add_break(WD_BREAK.PAGE)
        
        # Referencias Familiares - título en azul, negrita, mayúsculas, centrado
        # Solo agregar si hay referencias familiares
        if referencias_familiares:
            p_ref_fam_titulo = doc.add_paragraph()
            p_ref_fam_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ref_fam_titulo = p_ref_fam_titulo.add_run("REFERENCIAS FAMILIARES")
            run_ref_fam_titulo.bold = True
            run_ref_fam_titulo.font.size = Pt(12)
            run_ref_fam_titulo.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            doc.add_paragraph()
            
            for ref in referencias_familiares:
                nombre_ref = ref.get('nombre', '').strip()
                telefono_ref = ref.get('telefono', ref.get('celular', '')).strip()
                
                if nombre_ref:
                    p_ref_fam_nombre = doc.add_paragraph()
                    run_ref_fam_nombre = p_ref_fam_nombre.add_run(nombre_ref)
                    run_ref_fam_nombre.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_ref_fam_nombre.italic = True
                    
                    if telefono_ref:
                        p_ref_fam_tel = doc.add_paragraph()
                        run_ref_fam_tel_label = p_ref_fam_tel.add_run("Teléfono: ")
                        run_ref_fam_tel_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_ref_fam_tel_label.bold = True
                        run_ref_fam_tel_valor = p_ref_fam_tel.add_run(telefono_ref)
                        run_ref_fam_tel_valor.font.color.rgb = RGBColor(0, 0, 0)
                        run_ref_fam_tel_valor.bold = True
                    
                    doc.add_paragraph()
        
        # Referencias Personales - título en azul, negrita, mayúsculas, centrado
        # Solo agregar si hay referencias personales
        if referencias_personales:
            p_ref_per_titulo = doc.add_paragraph()
            p_ref_per_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ref_per_titulo = p_ref_per_titulo.add_run("REFERENCIAS PERSONALES")
            run_ref_per_titulo.bold = True
            run_ref_per_titulo.font.size = Pt(12)
            run_ref_per_titulo.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            doc.add_paragraph()
            
            for ref in referencias_personales:
                nombre_ref = ref.get('nombre', '').strip()
                telefono_ref = ref.get('telefono', ref.get('celular', '')).strip()
                
                if nombre_ref:
                    p_ref_per_nombre = doc.add_paragraph()
                    run_ref_per_nombre = p_ref_per_nombre.add_run(nombre_ref)
                    run_ref_per_nombre.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    run_ref_per_nombre.italic = True
                    
                    if telefono_ref:
                        p_ref_per_tel = doc.add_paragraph()
                        run_ref_per_tel_label = p_ref_per_tel.add_run("Teléfono: ")
                        run_ref_per_tel_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        run_ref_per_tel_label.bold = True
                        run_ref_per_tel_valor = p_ref_per_tel.add_run(telefono_ref)
                        run_ref_per_tel_valor.font.color.rgb = RGBColor(0, 0, 0)
                        run_ref_per_tel_valor.bold = True
                    
                    doc.add_paragraph()
        
        # Espacios finales antes del pie de página
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Pie de página con nombre en azul, negrita, mayúsculas
        p_final = doc.add_paragraph()
        run_final = p_final.add_run(nombre.upper())
        run_final.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run_final.bold = True
        
        p_cedula_final = doc.add_paragraph()
        run_cedula_final = p_cedula_final.add_run(f"C.C. {cedula} de {exp}")
        run_cedula_final.font.color.rgb = RGBColor(0, 0, 0)
        
        # Guardar en memoria
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Nombre del archivo
        nombre_archivo = nombre.replace(' ', '_') if nombre else 'Hoja_de_Vida'
        filename = f"HV_{nombre_archivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

def reemplazar_texto_en_documento(doc, reemplazos):
    """
    Reemplaza texto en un documento Word manteniendo el formato.
    Busca en párrafos y tablas. Busca placeholders de forma case-insensitive.
    Mejora: Busca en todos los runs de texto para encontrar variables divididas.
    """
    import re
    
    def reemplazar_en_parrafo(paragraph, reemplazos_dict):
        """Reemplaza texto en un párrafo manteniendo formato, especialmente negrilla"""
        # Obtener todo el texto del párrafo
        texto_original = paragraph.text
        if not texto_original:
            return
        
        # Trabajar directamente con los runs para conservar formato individual
        if not paragraph.runs:
            return
        
        # Buscar y reemplazar en cada run, conservando formato
        sorted_reemplazos = sorted(reemplazos_dict.items(), key=lambda x: len(x[0]), reverse=True)
        
        for placeholder, valor in sorted_reemplazos:
            if not placeholder:
                continue
            
            # Si el placeholder tiene formato {{VARIABLE}}, buscar solo en mayúsculas y case-sensitive
            if placeholder.startswith('{{') and placeholder.endswith('}}'):
                pattern = re.escape(placeholder)
                case_sensitive = True
            # Para dia1 y dia2, buscar sin word boundaries para capturar variaciones
            elif 'dia' in placeholder.lower() and len(placeholder) <= 5:
                pattern = re.escape(placeholder)
                case_sensitive = False
            elif placeholder.upper() in ['4 TURNOS', 'ADICIONALES', 'SUELDO FIJO MENSUAL', 'SUELDO PROPORCIONAL', 
                                         'BONO SEGURIDAD', 'AUXILIO DE TRANSPORTE', 'TURNOS', 'DESCANSOS']:
                continue
            elif 'SUELDO FIJO' in placeholder.upper() or 'SUELDO PROPORCIONAL' in placeholder.upper():
                continue
            elif 'BONO SEGURIDAD' in placeholder.upper() or 'AUXILIO' in placeholder.upper():
                continue
            else:
                pattern = re.escape(placeholder)
                case_sensitive = False
            
            # Buscar el placeholder en todos los runs
            texto_completo = ''.join([run.text for run in paragraph.runs])
            flags = 0 if case_sensitive else re.IGNORECASE
            matches = list(re.finditer(pattern, texto_completo, flags))
            
            if matches:
                # Reemplazar desde el final hacia el inicio para mantener índices
                for match in reversed(matches):
                    start, end = match.span()
                    valor_reemplazo = str(valor)
                    
                    # Para dia1 y dia2, verificar si necesita espacios alrededor
                    if 'dia' in placeholder.lower() and len(placeholder) <= 5:
                        tiene_espacio_antes = start > 0 and texto_completo[start-1].isspace()
                        tiene_espacio_despues = end < len(texto_completo) and texto_completo[end].isspace()
                        
                        if valor_reemplazo.startswith(' ') and tiene_espacio_antes:
                            valor_reemplazo = valor_reemplazo.lstrip()
                        if valor_reemplazo.endswith(' ') and tiene_espacio_despues:
                            valor_reemplazo = valor_reemplazo.rstrip()
                        
                        if not tiene_espacio_antes and not valor_reemplazo.startswith(' '):
                            if start > 0 and texto_completo[start-1].isalnum():
                                valor_reemplazo = ' ' + valor_reemplazo
                        if not tiene_espacio_despues and not valor_reemplazo.endswith(' '):
                            if end < len(texto_completo) and texto_completo[end].isalnum():
                                valor_reemplazo = valor_reemplazo + ' '
                    
                    # Encontrar en qué run(s) está el placeholder y conservar formato (especialmente negrilla)
                    current_pos = 0
                    formato_aplicar = None
                    placeholder_encontrado = False
                    
                    for run in paragraph.runs:
                        run_start = current_pos
                        run_end = current_pos + len(run.text)
                        
                        # Si el placeholder está completamente en este run
                        if run_start <= start < run_end and run_start <= end <= run_end:
                            # Conservar formato del run (incluyendo negrilla)
                            formato_aplicar = {
                                'font_name': run.font.name if run.font.name else None,
                                'font_size': run.font.size if run.font.size else None,
                                'bold': run.bold if run.bold is not None else False,
                                'italic': run.italic if run.italic is not None else False,
                                'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                            }
                            
                            # Reemplazar en el run - el formato se conserva automáticamente
                            run_start_in_run = start - run_start
                            run_end_in_run = end - run_start
                            nuevo_texto = run.text[:run_start_in_run] + valor_reemplazo + run.text[run_end_in_run:]
                            run.text = nuevo_texto
                            placeholder_encontrado = True
                            break
                        
                        # Si el placeholder comienza en este run (puede estar dividido)
                        elif run_start <= start < run_end and formato_aplicar is None:
                            # Usar el formato del run donde comienza el placeholder (conservar negrilla)
                            formato_aplicar = {
                                'font_name': run.font.name if run.font.name else None,
                                'font_size': run.font.size if run.font.size else None,
                                'bold': run.bold if run.bold is not None else False,
                                'italic': run.italic if run.italic is not None else False,
                                'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                            }
                        
                        current_pos = run_end
                    
                    # Si el placeholder estaba dividido entre múltiples runs, reconstruir conservando formato
                    if not placeholder_encontrado and formato_aplicar:
                        # Reconstruir texto completo
                        texto_completo_nuevo = ''.join([run.text for run in paragraph.runs])
                        if placeholder in texto_completo_nuevo:
                            # Reemplazar en el texto completo
                            texto_completo_nuevo = texto_completo_nuevo.replace(placeholder, valor_reemplazo, 1)
                            
                            # Limpiar y reconstruir con el formato del run donde estaba el placeholder
                            paragraph.clear()
                            nuevo_run = paragraph.add_run(texto_completo_nuevo)
                            
                            # Aplicar formato conservado (incluyendo negrilla)
                            if formato_aplicar:
                                if formato_aplicar['font_name']:
                                    nuevo_run.font.name = formato_aplicar['font_name']
                                if formato_aplicar['font_size']:
                                    nuevo_run.font.size = formato_aplicar['font_size']
                                if formato_aplicar['color']:
                                    nuevo_run.font.color.rgb = formato_aplicar['color']
                                nuevo_run.bold = formato_aplicar['bold']
                                nuevo_run.italic = formato_aplicar['italic']
    
    def reemplazar_en_runs(runs, reemplazos_dict):
        """Reemplaza texto en runs individuales conservando formato, especialmente negrilla"""
        texto_completo = ''.join([run.text for run in runs])
        if not texto_completo:
            return False
        
        texto_nuevo = texto_completo
        cambios_realizados = False
        formato_aplicar = None
        
        # Ordenar por longitud descendente
        sorted_reemplazos = sorted(reemplazos_dict.items(), key=lambda x: len(x[0]), reverse=True)
        
        for placeholder, valor in sorted_reemplazos:
            if not placeholder:
                continue
            
            # Si el placeholder tiene formato {{VARIABLE}}, buscar solo en mayúsculas y case-sensitive
            if placeholder.startswith('{{') and placeholder.endswith('}}'):
                pattern = re.escape(placeholder)
                matches = list(re.finditer(pattern, texto_nuevo))  # Case-sensitive
            else:
                pattern = re.escape(placeholder)
                matches = list(re.finditer(pattern, texto_nuevo, re.IGNORECASE))
            
            if matches:
                cambios_realizados = True
                # Encontrar el formato del run donde está el placeholder (conservar negrilla)
                for match in reversed(matches):
                    start, end = match.span()
                    
                    # Encontrar en qué run está el placeholder para conservar su formato
                    current_pos = 0
                    for run in runs:
                        run_start = current_pos
                        run_end = current_pos + len(run.text)
                        
                        if run_start <= start < run_end and formato_aplicar is None:
                            # Conservar formato del run donde está el placeholder (incluyendo negrilla)
                            formato_aplicar = {
                                'font_name': run.font.name if run.font.name else None,
                                'font_size': run.font.size if run.font.size else None,
                                'bold': run.bold if run.bold is not None else False,
                                'italic': run.italic if run.italic is not None else False,
                                'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                            }
                            break
                        
                        current_pos = run_end
                    
                    texto_nuevo = texto_nuevo[:start] + str(valor) + texto_nuevo[end:]
        
        if cambios_realizados and texto_nuevo != texto_completo:
            # Limpiar todos los runs y crear uno nuevo con el texto reemplazado
            for run in runs:
                run.text = ''
            if runs:
                runs[0].text = texto_nuevo
                # Aplicar formato conservado (incluyendo negrilla)
                if formato_aplicar:
                    if formato_aplicar['font_name']:
                        runs[0].font.name = formato_aplicar['font_name']
                    if formato_aplicar['font_size']:
                        runs[0].font.size = formato_aplicar['font_size']
                    if formato_aplicar['color']:
                        runs[0].font.color.rgb = formato_aplicar['color']
                    runs[0].bold = formato_aplicar['bold']
                    runs[0].italic = formato_aplicar['italic']
            return True
        
        return False
    
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        # Primero intentar reemplazo en el párrafo completo
        reemplazar_en_parrafo(paragraph, reemplazos)
        
        # También intentar reemplazo en runs individuales (por si la variable está dividida)
        if paragraph.runs and len(paragraph.runs) > 1:
            reemplazar_en_runs(paragraph.runs, reemplazos)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    reemplazar_en_parrafo(paragraph, reemplazos)
                    # También en runs individuales
                    if paragraph.runs and len(paragraph.runs) > 1:
                        reemplazar_en_runs(paragraph.runs, reemplazos)
    
    # Reemplazar en headers y footers de todas las secciones
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                reemplazar_en_parrafo(paragraph, reemplazos)
                if paragraph.runs and len(paragraph.runs) > 1:
                    reemplazar_en_runs(paragraph.runs, reemplazos)
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                reemplazar_en_parrafo(paragraph, reemplazos)
                if paragraph.runs and len(paragraph.runs) > 1:
                    reemplazar_en_runs(paragraph.runs, reemplazos)
        # First page header/footer (si es diferente)
        if section.different_first_page_header_footer:
            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    reemplazar_en_parrafo(paragraph, reemplazos)
                    if paragraph.runs and len(paragraph.runs) > 1:
                        reemplazar_en_runs(paragraph.runs, reemplazos)
            if section.first_page_footer:
                for paragraph in section.first_page_footer.paragraphs:
                    reemplazar_en_parrafo(paragraph, reemplazos)
                    if paragraph.runs and len(paragraph.runs) > 1:
                        reemplazar_en_runs(paragraph.runs, reemplazos)

def formatear_monto(monto, incluir_signo=True):
    """Formatea un monto como moneda colombiana"""
    if not monto:
        return ''
    try:
        monto_num = float(monto)
        # Formatear sin el símbolo $ para evitar duplicaciones en el template
        monto_formateado = f"{monto_num:,.0f}".replace(',', '.')
        if incluir_signo:
            return f"${monto_formateado}"
        return monto_formateado
    except:
        return str(monto)

def sanitize_input(value, max_length=500):
    """Sanitiza un input removiendo caracteres peligrosos y limitando longitud"""
    if not value:
        return ''
    # Remover caracteres de control y limitar longitud
    sanitized = ''.join(c for c in str(value) if c.isprintable() or c in ['\n', '\r', '\t'])[:max_length]
    return sanitized.strip()

def validate_numeric(value, min_val=None, max_val=None, default=0):
    """Valida y convierte un valor numérico"""
    try:
        num = float(str(value).replace(',', '.').replace('.', '', str(value).count('.') - 1))
        if min_val is not None and num < min_val:
            return default
        if max_val is not None and num > max_val:
            return default
        return num
    except (ValueError, AttributeError):
        return default

@app.route('/generate-cuenta-cobro', methods=['POST'])
def generate_cuenta_cobro():
    """Genera una cuenta de cobro usando el template Word"""
    try:
        if not request.json:
            return jsonify({'error': 'No se recibieron datos'}), 400
        
        data = request.json
        
        # Validar y sanitizar datos del formulario
        nombre = sanitize_input(data.get('nombre', ''), max_length=200)
        if not nombre:
            return jsonify({'error': 'El nombre es obligatorio'}), 400
        
        cedula = sanitize_input(data.get('cedula', ''), max_length=50)
        if not cedula:
            return jsonify({'error': 'La cédula es obligatoria'}), 400
        
        telefono = sanitize_input(data.get('phone', '') or data.get('telefono', '') or data.get('phoneNumber', ''), max_length=20)
        # Remover print de debug con datos sensibles en producción
        # print(f"📞 Teléfono recibido: '{telefono}'")  # Debug - removido por seguridad
        mes = sanitize_input(data.get('mes', ''), max_length=50)
        año = sanitize_input(data.get('año', ''), max_length=10)
        
        # Validar valores numéricos
        mes_completo = bool(data.get('mesCompleto', True))
        dia_inicio = str(int(validate_numeric(data.get('diaInicio', '1'), min_val=1, max_val=31, default=1)))
        dia_fin = str(int(validate_numeric(data.get('diaFin', '30'), min_val=1, max_val=31, default=30)))
        
        # Calcular días trabajados
        dias_num = 30  # Valor por defecto
        if not mes_completo:
            try:
                dia_inicio_num = int(dia_inicio)
                dia_fin_num = int(dia_fin)
                if dia_fin_num >= dia_inicio_num:
                    dias_num = (dia_fin_num - dia_inicio_num) + 1
                else:
                    dias_num = 30
            except (ValueError, TypeError):
                dias_num = 30
        else:
            # Si es mes completo, usar el valor del campo o calcular desde el mes
            try:
                dias_trabajados_input = data.get('diasTrabajados', '')
                if dias_trabajados_input:
                    dias_num = int(validate_numeric(dias_trabajados_input, min_val=1, max_val=31, default=30))
                else:
                    dias_num = 30
            except:
                dias_num = 30
        
        # Obtener el número de días del mes seleccionado
        try:
            mes_num = int(mes) if mes.isdigit() else 0
            año_num = int(año) if año.isdigit() else datetime.now().year
            if 1 <= mes_num <= 12:
                from calendar import monthrange
                dias_del_mes = monthrange(año_num, mes_num)[1]
            else:
                dias_del_mes = 30
        except:
            dias_del_mes = 30
        
        # Limitar días trabajados al máximo de días del mes
        if dias_num > dias_del_mes:
            dias_num = dias_del_mes
        if dias_num < 1:
            dias_num = 30
        
        # Parsear sueldo fijo correctamente (soporta formatos: "2000000", "2.000.000", "2,000,000")
        sueldo_fijo_num = 0
        try:
            sueldo_fijo_raw = str(data.get('sueldoFijo', '0')).strip()
            if sueldo_fijo_raw:
                # Remover puntos (separadores de miles) y reemplazar coma por punto (decimal)
                sueldo_fijo_limpio = sueldo_fijo_raw.replace('.', '').replace(',', '.')
                sueldo_fijo_num = float(sueldo_fijo_limpio)
                if sueldo_fijo_num < 0:
                    sueldo_fijo_num = 0
                if sueldo_fijo_num > 10000000:
                    sueldo_fijo_num = 10000000
        except (ValueError, TypeError):
            sueldo_fijo_num = 0
        
        # Calcular sueldo proporcional según días trabajados
        sueldo_proporcional = 0
        if sueldo_fijo_num > 0 and dias_del_mes > 0:
            valor_por_dia = sueldo_fijo_num / dias_del_mes
            sueldo_proporcional = round(valor_por_dia * dias_num)
        
        # Validar y sanitizar valores monetarios y otros campos
        turnos_descansos = str(int(validate_numeric(data.get('turnosDescansos', '0'), min_val=0, max_val=100, default=0)))
        paciente = sanitize_input(data.get('paciente', '') or data.get('patientName', ''), max_length=200)
        cuenta_bancaria = sanitize_input(data.get('cuentaBancaria', '') or data.get('bankAccount', ''), max_length=50)
        banco = sanitize_input(data.get('banco', ''), max_length=100).upper() or 'Bancolombia'
        tipo_cuenta_cobro = sanitize_input(data.get('tipoCuentaCobro', '12h'), max_length=10)
        if tipo_cuenta_cobro not in ['12h', '8h']:
            tipo_cuenta_cobro = '12h'
        tiene_auxilio_transporte = bool(data.get('tieneAuxilioTransporte', False))
        
        # Parsear bono de seguridad CORRECTAMENTE
        bono_seguridad_num = 0
        try:
            bono_raw = data.get('bonoSeguridad', '0')
            if bono_raw:
                # Convertir a string y limpiar
                bono_str = str(bono_raw).strip()
                # Remover puntos (separadores de miles) y reemplazar coma por punto (decimal)
                bono_limpio = bono_str.replace('.', '').replace(',', '.')
                bono_seguridad_num = float(bono_limpio)
                # Validar rango
                if bono_seguridad_num < 0:
                    bono_seguridad_num = 0
                if bono_seguridad_num > 10000000:
                    bono_seguridad_num = 10000000
        except (ValueError, TypeError, AttributeError):
            bono_seguridad_num = 0
        
        # Parsear auxilio de transporte
        auxilio_transporte_num = 0
        if tiene_auxilio_transporte:
            try:
                auxilio_raw = data.get('auxilioTransporte', '0')
                if auxilio_raw:
                    auxilio_str = str(auxilio_raw).strip()
                    auxilio_limpio = auxilio_str.replace('.', '').replace(',', '.')
                    auxilio_transporte_num = float(auxilio_limpio)
                    if auxilio_transporte_num < 0:
                        auxilio_transporte_num = 0
                    if auxilio_transporte_num > 10000000:
                        auxilio_transporte_num = 10000000
            except (ValueError, TypeError, AttributeError):
                auxilio_transporte_num = 0
        
        # Calcular adicionales (turnos * 60000)
        turnos_num = int(turnos_descansos) if turnos_descansos.isdigit() else 0
        valor_por_turno = 60000
        adicionales_valor = turnos_num * valor_por_turno
        
        # El total se calculará después de formatear los valores
        
        # Formatear fecha (mes en texto)
        fecha_texto = ''
        if mes and año:
            mes_num = int(mes) if mes.isdigit() else 0
            if 1 <= mes_num <= 12:
                fecha_texto = f"{MESES[mes_num].upper()} DE {año}"
        
        # Cargar template
        # Seleccionar template según tipo de cuenta de cobro
        templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
        if tipo_cuenta_cobro == '8h':
            template_path = os.path.join(templates_dir, 'cobro_8h.docx')
        else:
            # Probar ambos nombres (con y sin espacio - typo común)
            for name in ['cobro_2026.docx', 'cobro_ 2026.docx']:
                p = os.path.join(templates_dir, name)
                if os.path.exists(p):
                    template_path = p
                    break
            else:
                template_path = os.path.join(templates_dir, 'cobro_2026.docx')
        if not os.path.exists(template_path):
            return jsonify({"error": f"Template no encontrado en: {template_path}"}), 404
        
        doc = Document(template_path)
        
        # Preparar reemplazos usando los placeholders exactos del template
        # Buscar todas las variaciones posibles de las variables
        reemplazos = {}
        
        # ============================================
        # FORMATEO DE VALORES - CÓDIGO NUEVO DESDE CERO
        # ============================================
        
        # Variable sf1: Sueldo proporcional según días trabajados
        # Ejemplo: sueldo fijo 2.000.000, enero 31 días = 2.000.000, febrero 28 días = 2.000.000
        sf1_valor = sueldo_proporcional
        sf1_formateado = formatear_monto(sf1_valor, incluir_signo=False)
        
        # Variable bs1: Bono de seguridad
        # Ejemplo: 200.000
        bs1_valor = bono_seguridad_num
        bs1_formateado = formatear_monto(bs1_valor, incluir_signo=False) if bs1_valor > 0 else ''
        
        # Variable ad1: Adicionales (turnos de descansos)
        # Ejemplo: 4 turnos * 60.000 = 240.000
        ad1_valor = adicionales_valor
        ad1_formateado = formatear_monto(ad1_valor, incluir_signo=False) if ad1_valor > 0 else ''
        
        # Variable ax1: Auxilio de transporte
        ax1_valor = auxilio_transporte_num
        ax1_formateado = formatear_monto(ax1_valor, incluir_signo=False) if ax1_valor > 0 else ''
        
        # Calcular TOTAL: sf1 + bs1 + ad1 + ax1
        total_calculado = sf1_valor + bs1_valor + ad1_valor + ax1_valor
        total_formateado = formatear_monto(total_calculado, incluir_signo=False)
        
        # SOLO reemplazar variables {{VARIABLE}} - no tocar texto normal
        # El template Word usa placeholders entre llaves dobles
        nombre_upper = nombre.upper()
        paciente_upper = paciente.upper() if paciente else ''
        telefono_valor = telefono if telefono else ''
        
        # Variables del template (solo formato {{variable}})
        reemplazos['{{Name1}}'] = nombre_upper
        reemplazos['{{Cedu1}}'] = cedula
        reemplazos['{{Num1}}'] = telefono_valor
        reemplazos['{{banco1}}'] = banco
        reemplazos['{{nbanco1}}'] = cuenta_bancaria
        reemplazos['{{mes1}}'] = fecha_texto
        reemplazos['{{valor1}}'] = total_formateado
        reemplazos['{{paciente1}}'] = paciente_upper
        reemplazos['{{sf1}}'] = sf1_formateado
        reemplazos['{{sb1}}'] = bs1_formateado
        reemplazos['{{ad1}}'] = ad1_formateado
        reemplazos['{{ax1}}'] = ax1_formateado
        
        # Días trabajados - {{dias1}} para uso numérico; MES COMPLETO se deja tal cual (no reemplazar con días)
        reemplazos['{{dias1}}'] = str(dias_num)
        
        # dia1 y dia2 - solo {{dia1}} {{dia2}}
        dia_inicio = data.get('diaInicio', '1').strip() if data.get('diaInicio') else '1'
        dia_fin = str(dias_del_mes)
        reemplazos['{{dia1}}'] = dia_inicio
        reemplazos['{{dia2}}'] = dia_fin
        
        # Limpiar duplicaciones de texto comunes ANTES de reemplazar
        # Duplicaciones de año - múltiples variaciones (ordenar por longitud descendente)
        # Primero los más largos para evitar reemplazos parciales
        reemplazos['DE ' + año + ' DEL ' + año] = f'DE {año}'
        reemplazos['DEL ' + año + ' DE ' + año] = f'DEL {año}'
        reemplazos['DE ' + año + ' DE ' + año] = f'DE {año}'
        reemplazos['DEL ' + año + ' DEL ' + año] = f'DEL {año}'
        # También valores hardcodeados comunes
        reemplazos['DE 2026 DEL 2026'] = f'DE {año}'
        reemplazos['DEL 2026 DE 2026'] = f'DEL {año}'
        reemplazos['DE 2026 DE 2026'] = f'DE {año}'
        reemplazos['DEL 2026 DEL 2026'] = f'DEL {año}'
        
        # Log de reemplazos para debug - especialmente dia1 y dia2
        print(f"🔍 Reemplazos a realizar: {len(reemplazos)} variables")
        print(f"📅 dia1 (día inicio): '{dia_inicio}'")
        print(f"📅 dia2 (día fin): '{dia_fin}'")
        # Debug removido por seguridad - comentado para producción
        # for key, value in sorted(reemplazos.items()):
        #     if value and ('dia' in key.lower() or 'DIA' in key):
        #         print(f"  - {key} -> {value}")
        
        # Reemplazar texto en el documento
        reemplazar_texto_en_documento(doc, reemplazos)
        print("✅ Reemplazos completados en el documento")
        
        # Verificar si dia1 y dia2 fueron reemplazados correctamente
        texto_completo = ' '.join([para.text for para in doc.paragraphs])
        if 'dia1' in texto_completo.lower() or 'dia2' in texto_completo.lower():
            # Debug removido por seguridad - solo en desarrollo
            # print(f"⚠️ ADVERTENCIA: Todavía hay 'dia1' o 'dia2' sin reemplazar en el documento")
            # print(f"   Texto encontrado: {texto_completo[texto_completo.lower().find('dia'):texto_completo.lower().find('dia')+50]}")
            pass
        
        # Limpiar duplicaciones después del reemplazo
        # Buscar y limpiar patrones comunes de duplicación
        import re
        for paragraph in doc.paragraphs:
            texto = paragraph.text
            # Limpiar duplicaciones de año - múltiples patrones (aplicar varias veces para asegurar)
            # Patrón 1: DE 2026 DE 2026 -> DE 2026
            texto = re.sub(r'DE (\d{4}) DE \1', r'DE \1', texto)
            # Patrón 2: DEL 2026 DEL 2026 -> DEL 2026
            texto = re.sub(r'DEL (\d{4}) DEL \1', r'DEL \1', texto)
            # Patrón 3: DE 2026 DEL 2026 -> DE 2026 (el más común)
            texto = re.sub(r'DE (\d{4}) DEL \1', r'DE \1', texto)
            # Patrón 4: DEL 2026 DE 2026 -> DEL 2026
            texto = re.sub(r'DEL (\d{4}) DE \1', r'DEL \1', texto)
            # Aplicar nuevamente para casos anidados
            texto = re.sub(r'DE (\d{4}) DEL \1', r'DE \1', texto)
            texto = re.sub(r'DEL (\d{4}) DE \1', r'DEL \1', texto)
            # Limpiar múltiples símbolos $ seguidos
            texto = re.sub(r'\$\$+', '$', texto)
            texto = re.sub(r'\$ \$+', '$', texto)
            # Limpiar espacios múltiples
            texto = re.sub(r'  +', ' ', texto)
            
            if texto != paragraph.text:
                # Guardar formato
                formato_original = None
                if paragraph.runs:
                    primer_run = paragraph.runs[0]
                    formato_original = {
                        'font_name': primer_run.font.name if primer_run.font.name else None,
                        'font_size': primer_run.font.size if primer_run.font.size else None,
                        'bold': primer_run.bold if primer_run.bold is not None else False,
                        'italic': primer_run.italic if primer_run.italic is not None else False,
                        'color': primer_run.font.color.rgb if primer_run.font.color and primer_run.font.color.rgb else None
                    }
                
                paragraph.clear()
                nuevo_run = paragraph.add_run(texto)
                
                if formato_original:
                    if formato_original['font_name']:
                        nuevo_run.font.name = formato_original['font_name']
                    if formato_original['font_size']:
                        nuevo_run.font.size = formato_original['font_size']
                    if formato_original['color']:
                        nuevo_run.font.color.rgb = formato_original['color']
                    nuevo_run.bold = formato_original['bold']
                    nuevo_run.italic = formato_original['italic']
        
        # También limpiar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        texto = paragraph.text
                        # Limpiar duplicaciones de año - múltiples patrones (aplicar varias veces)
                        # Patrón 1: DE 2026 DE 2026 -> DE 2026
                        texto = re.sub(r'DE (\d{4}) DE \1', r'DE \1', texto)
                        # Patrón 2: DEL 2026 DEL 2026 -> DEL 2026
                        texto = re.sub(r'DEL (\d{4}) DEL \1', r'DEL \1', texto)
                        # Patrón 3: DE 2026 DEL 2026 -> DE 2026 (el más común)
                        texto = re.sub(r'DE (\d{4}) DEL \1', r'DE \1', texto)
                        # Patrón 4: DEL 2026 DE 2026 -> DEL 2026
                        texto = re.sub(r'DEL (\d{4}) DE \1', r'DEL \1', texto)
                        # Aplicar nuevamente para casos anidados
                        texto = re.sub(r'DE (\d{4}) DEL \1', r'DE \1', texto)
                        texto = re.sub(r'DEL (\d{4}) DE \1', r'DEL \1', texto)
                        texto = re.sub(r'\$\$+', '$', texto)
                        texto = re.sub(r'\$ \$+', '$', texto)
                        texto = re.sub(r'  +', ' ', texto)
                        
                        if texto != paragraph.text:
                            formato_original = None
                            if paragraph.runs:
                                primer_run = paragraph.runs[0]
                                formato_original = {
                                    'font_name': primer_run.font.name if primer_run.font.name else None,
                                    'font_size': primer_run.font.size if primer_run.font.size else None,
                                    'bold': primer_run.bold if primer_run.bold is not None else False,
                                    'italic': primer_run.italic if primer_run.italic is not None else False,
                                    'color': primer_run.font.color.rgb if primer_run.font.color and primer_run.font.color.rgb else None
                                }
                            
                            paragraph.clear()
                            nuevo_run = paragraph.add_run(texto)
                            
                            if formato_original:
                                if formato_original['font_name']:
                                    nuevo_run.font.name = formato_original['font_name']
                                if formato_original['font_size']:
                                    nuevo_run.font.size = formato_original['font_size']
                                if formato_original['color']:
                                    nuevo_run.font.color.rgb = formato_original['color']
                                nuevo_run.bold = formato_original['bold']
                                nuevo_run.italic = formato_original['italic']
        
        print("✅ Limpieza de duplicaciones completada")
        
        # Procesar tablas: eliminar fila de ADICIONALES si no hay turnos, agregar AUXILIO DE TRANSPORTE si está seleccionado
        for table in doc.tables:
            filas_a_eliminar = []
            indice_adicionales = -1
            indice_ultima_fila_datos = -1
            
            # Buscar la fila de ADICIONALES y la última fila de datos (antes del TOTAL)
            for idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.strip() for cell in row.cells]).upper()
                
                # Buscar fila de ADICIONALES
                if 'ADICIONALES' in row_text and 'SUELDO FIJO' not in row_text and 'BONO' not in row_text:
                    indice_adicionales = idx
                
                # Buscar última fila de datos (antes del TOTAL)
                if 'TOTAL' not in row_text and 'SUELDO FIJO' in row_text or 'BONO' in row_text or 'ADICIONALES' in row_text:
                    indice_ultima_fila_datos = idx
            
            # Eliminar fila de ADICIONALES si no hay turnos
            if turnos_num == 0 and indice_adicionales >= 0:
                filas_a_eliminar.append(indice_adicionales)
                print(f"🗑️ Eliminando fila de ADICIONALES (índice {indice_adicionales}) - no hay turnos")
            
            # Eliminar filas en orden inverso para mantener los índices correctos
            for idx in sorted(filas_a_eliminar, reverse=True):
                if idx < len(table.rows):
                    table._element.remove(table.rows[idx]._element)
            
            # Agregar fila de AUXILIO DE TRANSPORTE si está seleccionado
            if tiene_auxilio_transporte and auxilio_transporte_num > 0:
                # Buscar la fila del TOTAL para insertar antes de ella
                indice_total = -1
                for idx, row in enumerate(table.rows):
                    row_text = ' '.join([cell.text.strip() for cell in row.cells]).upper()
                    if 'TOTAL' in row_text:
                        indice_total = idx
                        break
                
                # Si no se encuentra TOTAL, usar el final de la tabla
                if indice_total < 0:
                    indice_total = len(table.rows)
                
                # Crear nueva fila al final primero
                nueva_fila = table.add_row()
                
                # Llenar las celdas de la nueva fila
                if len(nueva_fila.cells) >= 4:
                    # Columna 1: Descripción
                    nueva_fila.cells[0].text = 'AUXILIO DE TRANSPORTE'
                    # Columna 2: Cantidad
                    nueva_fila.cells[1].text = 'MES COMPLETO'
                    # Columna 3: Valor
                    nueva_fila.cells[2].text = ax1_formateado
                    # Columna 4: Paciente (vacía)
                    nueva_fila.cells[3].text = ''
                
                # Mover la fila a la posición correcta (antes del TOTAL)
                if indice_total < len(table.rows) - 1:
                    nueva_fila_element = nueva_fila._element
                    tbl = table._element
                    # Remover de la posición actual
                    tbl.remove(nueva_fila_element)
                    # Insertar antes del TOTAL
                    fila_total = table.rows[indice_total]._element
                    fila_total.addprevious(nueva_fila_element)
                
                print(f"✅ Agregada fila de AUXILIO DE TRANSPORTE con valor {ax1_formateado}")
        
        # Guardar en memoria
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Nombre del archivo
        nombre_archivo = nombre.replace(' ', '_') if nombre else 'Cuenta_Cobro'
        filename = f"Cuenta_Cobro_{nombre_archivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/generate-contrato-arrendamiento', methods=['POST'])
def generate_contrato_arrendamiento():
    """Genera un contrato de arrendamiento de predio rural usando el template Word"""
    try:
        if not request.json:
            return jsonify({'error': 'No se recibieron datos'}), 400
        
        data = request.json
        
        # Validar y sanitizar datos del formulario
        nombre_arrendador = sanitize_input(data.get('nombreArrendador', ''), max_length=200)
        if not nombre_arrendador:
            return jsonify({'error': 'El nombre del arrendador es obligatorio'}), 400
        
        cedula_arrendador = sanitize_input(data.get('cedulaArrendador', ''), max_length=50)
        if not cedula_arrendador:
            return jsonify({'error': 'La cédula del arrendador es obligatoria'}), 400
        
        ciudad_expedicion_arrendador = sanitize_input(data.get('ciudadExpedicionArrendador', ''), max_length=100)
        nombre_arrendatario = sanitize_input(data.get('nombreArrendatario', ''), max_length=200)
        cedula_arrendatario = sanitize_input(data.get('cedulaArrendatario', ''), max_length=50)
        ciudad_expedicion_arrendatario = sanitize_input(data.get('ciudadExpedicionArrendatario', ''), max_length=100)
        nombre_predio = sanitize_input(data.get('nombrePredio', ''), max_length=200)
        nombre_vereda = sanitize_input(data.get('nombreVereda', ''), max_length=200)
        municipio = sanitize_input(data.get('municipio', ''), max_length=100)
        departamento = sanitize_input(data.get('departamento', ''), max_length=100)
        direccion_referencia = sanitize_input(data.get('direccionReferencia', ''), max_length=500)
        hectareas_arrendadas = sanitize_input(data.get('hectareasArrendadas', ''), max_length=50)
        hectareas_totales = sanitize_input(data.get('hectareasTotales', ''), max_length=50)
        valor_canon = sanitize_input(data.get('valorCanon', ''), max_length=50)
        duracion_contrato_anios = sanitize_input(data.get('duracionContratoAnios', ''), max_length=10)
        fecha_inicio_contrato = sanitize_input(data.get('fechaInicioContrato', ''), max_length=50)
        ciudad_firma_contrato = sanitize_input(data.get('ciudadFirmaContrato', ''), max_length=100)
        dia_firma = sanitize_input(data.get('diaFirma', ''), max_length=10)
        mes_firma = sanitize_input(data.get('mesFirma', ''), max_length=10)
        anio_firma = sanitize_input(data.get('anioFirma', ''), max_length=10)
        
        # Obtener hectáreas en texto si viene del formulario
        hectareas_arrendadas_texto = data.get('hectareasArrendadasTexto', '')
        if not hectareas_arrendadas_texto and hectareas_arrendadas:
            try:
                hectareas_num = float(hectareas_arrendadas.replace(',', '.'))
                # Convertir a texto simple
                hectareas_arrendadas_texto = str(hectareas_num)
            except:
                hectareas_arrendadas_texto = hectareas_arrendadas
        
        # Obtener nombre del mes
        mes_nombre = data.get('mesFirmaNombre', '')
        if not mes_nombre and mes_firma:
            try:
                mes_num = int(mes_firma)
                if 1 <= mes_num <= 12:
                    mes_nombre = MESES.get(mes_num, mes_firma)
            except:
                mes_nombre = mes_firma
        
        # Cargar template
        base_dir = os.path.dirname(__file__)
        template_path = os.path.join(base_dir, 'templates', 'contrato.docx')
        
        # Debug: Listar archivos en templates si no existe
        if not os.path.exists(template_path):
            templates_dir = os.path.join(base_dir, 'templates')
            available_files = []
            if os.path.exists(templates_dir):
                available_files = os.listdir(templates_dir)
            error_msg = f"Template no encontrado en: {template_path}\n"
            error_msg += f"Directorio base: {base_dir}\n"
            error_msg += f"Directorio templates: {templates_dir}\n"
            error_msg += f"Archivos disponibles en templates: {', '.join(available_files) if available_files else 'Ninguno'}"
            return jsonify({"error": error_msg}), 404
        
        doc = Document(template_path)
        
        # Preparar reemplazos con todas las variaciones posibles
        reemplazos = {}
        
        # Solo reemplazar variables con formato {{VARIABLE}} (llaves dobles y mayúsculas)
        # Arrendador
        reemplazos['{{NOMBRE_ARRENDADOR}}'] = nombre_arrendador.upper()
        reemplazos['{{CEDULA_ARRENDADOR}}'] = cedula_arrendador
        reemplazos['{{CIUDAD_EXPEDICION_ARRENDADOR}}'] = ciudad_expedicion_arrendador.upper()
        
        # Arrendatario
        reemplazos['{{NOMBRE_ARRENDATARIO}}'] = nombre_arrendatario.upper()
        reemplazos['{{CEDULA_ARRENDATARIO}}'] = cedula_arrendatario
        reemplazos['{{CIUDAD_EXPEDICION_ARRENDATARIO}}'] = ciudad_expedicion_arrendatario.upper()
        
        # Predio
        reemplazos['{{NOMBRE_PREDIO}}'] = nombre_predio.upper()
        reemplazos['{{NOMBRE_VEREDA}}'] = nombre_vereda.upper()
        reemplazos['{{MUNICIPIO}}'] = municipio.upper()
        reemplazos['{{DEPARTAMENTO}}'] = departamento.upper()
        reemplazos['{{DIRECCION_REFERENCIA}}'] = direccion_referencia.upper()
        
        # Hectáreas
        reemplazos['{{HECTAREAS_ARRENDADAS}}'] = hectareas_arrendadas
        reemplazos['{{HECTAREAS_ARRENDADAS_TEXTO}}'] = hectareas_arrendadas_texto.upper()
        reemplazos['{{HECTAREAS_TOTALES}}'] = hectareas_totales
        
        # Valor del canon
        reemplazos['{{VALOR_CANON}}'] = valor_canon
        
        # Duración y fecha inicio
        reemplazos['{{DURACION_CONTRATO_ANIOS}}'] = duracion_contrato_anios
        
        # Formatear fecha de inicio
        fecha_inicio_formateada = ''
        if fecha_inicio_contrato:
            try:
                fecha_obj = datetime.strptime(fecha_inicio_contrato, '%Y-%m-%d')
                fecha_inicio_formateada = formatear_fecha(fecha_inicio_contrato)
            except:
                fecha_inicio_formateada = fecha_inicio_contrato
        
        reemplazos['{{FECHA_INICIO_CONTRATO}}'] = fecha_inicio_formateada
        
        # Firma
        reemplazos['{{CIUDAD_FIRMA_CONTRATO}}'] = ciudad_firma_contrato.upper()
        reemplazos['{{DIA_FIRMA}}'] = dia_firma
        reemplazos['{{MES_FIRMA}}'] = mes_nombre.upper()
        reemplazos['{{ANIO_FIRMA}}'] = anio_firma
        
        # Reemplazar texto en el documento
        reemplazar_texto_en_documento(doc, reemplazos)
        
        # Limpiar duplicaciones después del reemplazo
        # Buscar y limpiar patrones comunes de duplicación
        import re
        for paragraph in doc.paragraphs:
            texto = paragraph.text
            # Limpiar duplicaciones específicas primero
            # "CONVENCIÓN de CONVENCIÓN" -> "CONVENCIÓN"
            texto = re.sub(r'\b(CONVENCIÓN)\s+de\s+\1\b', r'\1', texto, flags=re.IGNORECASE)
            # "NORTE DE SANTANDER de NORTE DE SANTANDER" -> "NORTE DE SANTANDER"
            texto = re.sub(r'\b(NORTE DE SANTANDER)\s+de\s+\1\b', r'\1', texto, flags=re.IGNORECASE)
            # Limpiar duplicaciones generales: "TEXTO de TEXTO" -> "TEXTO"
            # Aplicar múltiples veces para casos anidados
            for _ in range(3):  # Aplicar hasta 3 veces para casos complejos
                # Patrón que captura cualquier texto seguido de " de " y el mismo texto
                texto_anterior = texto
                # Mejorar el patrón para capturar mejor textos con espacios
                texto = re.sub(r'\b([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{2,}?)\s+de\s+\1\b', r'\1', texto, flags=re.IGNORECASE)
                if texto == texto_anterior:
                    break  # No hay más cambios
            # Limpiar duplicaciones de año - múltiples patrones
            texto = re.sub(r'DE (\d{4}) DE \1', r'DE \1', texto)
            texto = re.sub(r'DEL (\d{4}) DEL \1', r'DEL \1', texto)
            texto = re.sub(r'DE (\d{4}) DEL \1', r'DE \1', texto)
            texto = re.sub(r'DEL (\d{4}) DE \1', r'DEL \1', texto)
            # Aplicar nuevamente para casos anidados
            texto = re.sub(r'DE (\d{4}) DEL \1', r'DE \1', texto)
            texto = re.sub(r'DEL (\d{4}) DE \1', r'DEL \1', texto)
            # Limpiar espacios múltiples
            texto = re.sub(r'  +', ' ', texto)
            
            if texto != paragraph.text:
                # Limpiar el párrafo y reconstruirlo
                paragraph.clear()
                paragraph.add_run(texto)
        
        # Guardar en memoria
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Nombre del archivo
        nombre_archivo = nombre_arrendador.replace(' ', '_') if nombre_arrendador else 'Contrato_Arrendamiento'
        filename = f"Contrato_Arrendamiento_{nombre_archivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

if __name__ == '__main__':
    # Crear directorio de templates si no existe
    os.makedirs(os.path.join(os.path.dirname(__file__), 'templates'), exist_ok=True)
    app.run(debug=True, host='0.0.0.0', port=5000)
